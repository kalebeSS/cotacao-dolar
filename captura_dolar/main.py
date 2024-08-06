from selenium import webdriver
from bs4 import BeautifulSoup
import requests
import pyautogui
import time
import datetime as dt
import re
from docx import Document
from docx.shared import Cm

def print_site():
    driver = webdriver.Chrome()
    driver.get('https://www.google.com/finance/quote/USD-BRL')
    driver.maximize_window()
    time.sleep(5)
    sc = pyautogui.screenshot()
    sc.save('cotacao.png')
    print('print salva')

def pegar_cotacao():
    response = requests.get('https://www.google.com/finance/quote/USD-BRL')
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        cotacao = soup.find('div', class_ = re.compile('fxKbKc'))
        valor = float(cotacao.text)
        return round(valor, 2)

    else:
        return 'requisição não processada'

def cria_documento():
    doc = Document()
    data = dt.date.today()
    site = 'https://www.google.com/finance/quote/USD-BRL'
    dolar = pegar_cotacao()
    doc.add_heading(f'Cotação Atual do Dólar – {dolar} ({data})')
    doc.add_paragraph(f'O Dólar está no valor de R${dolar}, na data {data}')
    doc.add_paragraph(f'Valor cotado no site {site}')
    doc.add_paragraph('Print da cotação')
    doc.add_picture('cotacao.png', width=Cm(15))
    doc.add_paragraph('Cotação feita por Kalebe')
    doc.save('dolar cotacao.docx')
    print('Documento criado')

def main():
    print_site()
    time.sleep(2)
    cria_documento()

main()